import docx

doc = docx.Document('/Users/parikshithkulkarni/Downloads/demo.docx')

for para in doc.paragraphs:
    print(para.text)
    for run in para.runs:
        print(run.text)

doc.add_paragraph('This is a new paragraph')


doc.save('newDoc.docx')


